# src/interfaces/telegram/utils/document.py

from datetime import datetime
from typing import Any, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


def create_search_results_doc(patents: List[Any], summaries: Optional[List[dict]] = None) -> str:
    """Создает DOC файл с результатами поиска"""
    doc = Document()
    
    # Устанавливаем альбомную ориентацию страницы
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)  # A4 ширина
    section.page_height = Inches(8.27)  # A4 высота
    
    # Устанавливаем поля страницы
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Настраиваем стиль документа
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)  # Уменьшаем размер шрифта
    
    # Добавляем заголовок
    title = doc.add_paragraph('Результаты поиска патентов')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(14)
    
    # Добавляем дату создания
    date_paragraph = doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Задаем ширину столбцов (в процентах от ширины страницы)
    widths = [5, 12, 25, 12, 12, 14, 14, 6]  # Проценты
    total_width = sum(widths)
    page_width = section.page_width - section.left_margin - section.right_margin
    
    for i, width in enumerate(widths):
        table.columns[i].width = int(page_width * width / total_width)
    
    # Заголовки столбцов
    headers = [
        '№', 'Номер патента', 'Название', 'Дата публикации', 'Дата заявки',
        'Авторы', 'Патентообладатели', 'МПК'
    ]
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Заполняем таблицу данными
    for idx, patent in enumerate(patents, 1):
        # Основная строка с информацией о патенте
        row_cells = table.add_row().cells
        
        # Форматируем даты
        pub_date = patent.publication_date.strftime("%d.%m.%Y") if patent.publication_date else "Не указана"
        app_date = patent.application_date.strftime("%d.%m.%Y") if patent.application_date else "Не указана"
        
        # Заполняем ячейки
        cells_data = [
            (0, str(idx)),
            (1, patent.id),
            (2, patent.title),
            (3, pub_date),
            (4, app_date),
            (5, ', '.join(patent.authors)),
            (6, ', '.join(patent.patent_holders)),
            (7, ', '.join(patent.ipc_codes))
        ]
        
        for col_idx, text in cells_data:
            cell = row_cells[col_idx]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Добавляем строку с описанием
        desc_row = table.add_row()
        
        # Объединяем ячейку с номером по вертикали
        desc_row.cells[0].merge(row_cells[0])
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Центрируем номер
        
        # Объединяем остальные ячейки для описания
        desc_cell = desc_row.cells[1]
        for i in range(2, len(headers)):
            desc_cell.merge(desc_row.cells[i])
        
        # Формируем текст описания
        if summaries and idx-1 < len(summaries) and summaries[idx-1] and summaries[idx-1].get("status") == "success":
            summary = summaries[idx-1]["summary"]
            if isinstance(summary, dict):
                desc_text = []
                if summary.get('description'):
                    desc_text.append(summary['description'])
                if summary.get('advantages'):
                    desc_text.append("Преимущества:\n• " + "\n• ".join(summary['advantages']))
                if summary.get('disadvantages'):
                    desc_text.append("Недостатки:\n• " + "\n• ".join(summary['disadvantages']))
                if summary.get('applications'):
                    desc_text.append("Области применения:\n• " + "\n• ".join(summary['applications']))
                desc_cell.text = "\n\n".join(desc_text)
            else:
                desc_cell.text = str(summary)
        else:
            desc_cell.text = patent.abstract if patent.abstract else "Нет данных"
        
        # Добавляем отступ для визуального разделения патентов
        desc_cell.paragraphs[0].paragraph_format.space_after = Pt(12)
        
        # Устанавливаем светло-серый фон для строки с описанием
        for paragraph in desc_cell.paragraphs:
            paragraph.style = doc.styles['Normal']
            paragraph.paragraph_format.first_line_indent = Inches(0.2)
    
    # Применяем автоподбор высоты строк
    for row in table.rows:
        row.height_rule = None
    
    # Сохраняем документ
    filename = f'search_results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    return filename